import os
import tempfile
import base64
import hashlib
import zipfile
import traceback
import re
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.background import BackgroundTask
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import requests
import shutil
from io import BytesIO

app = FastAPI(title="Word to Markdown Converter", version="2.1")

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# 从环境变量读取配置
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_REPO = os.getenv("GITHUB_REPO")

# ========== 辅助函数 ==========
def upload_image_to_github(image_data: bytes, image_name: str) -> str:
    try:
        hash_name = hashlib.md5(image_data).hexdigest()[:12]
        safe_name = f"{hash_name}_{image_name}"
        content = base64.b64encode(image_data).decode("utf-8")
        url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/images/{safe_name}"
        headers = {
            "Authorization": f"token {GITHUB_TOKEN}",
            "Accept": "application/vnd.github.v3+json"
        }
        
        check_resp = requests.get(url, headers=headers)
        if check_resp.status_code == 200:
            username, repo = GITHUB_REPO.split("/")
            return f"https://cdn.jsdelivr.net/gh/{username}/{repo}@main/images/{safe_name}"
        
        data = {
            "message": f"Upload {safe_name} via word2md",
            "content": content,
            "branch": "main"
        }
        resp = requests.put(url, json=data, headers=headers)
        
        if resp.status_code in [200, 201]:
            username, repo = GITHUB_REPO.split("/")
            return f"https://cdn.jsdelivr.net/gh/{username}/{repo}@main/images/{safe_name}"
        raise Exception(f"GitHub API returned {resp.status_code}")
    except Exception as e:
        print(f"Upload error: {str(e)}")
        raise

def extract_images_from_docx(docx_path: str) -> dict:
    """提取 docx 中的所有图片，返回 {图片名: 图片数据}"""
    images = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            for file_info in docx_zip.filelist:
                if file_info.filename.startswith('word/media/') and not file_info.filename.endswith('/'):
                    image_name = os.path.basename(file_info.filename)
                    images[image_name] = docx_zip.read(file_info.filename)
    except Exception as e:
        print(f"Error extracting images: {e}")
    return images

def get_paragraph_images(paragraph) -> list:
    """获取段落中的图片关系 ID 列表"""
    images = []
    try:
        # 查找所有内联图片
        for run in paragraph.runs:
            # 查找 blip 元素（图片）
            blips = run._element.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed_id = blip.get(qn('r:embed'))
                if embed_id:
                    images.append(embed_id)
    except Exception as e:
        print(f"Error getting paragraph images: {e}")
    return images

def get_image_relation_mapping(docx_path: str) -> dict:
    """获取图片关系 ID 到文件名的映射"""
    mapping = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 读取关系文件
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path in docx_zip.namelist():
                rels_content = docx_zip.read(rels_path).decode('utf-8')
                # 解析 XML 找到图片关系
                import xml.etree.ElementTree as ET
                root = ET.fromstring(rels_content)
                for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    rel_id = rel.get('Id')
                    target = rel.get('Target')
                    if target and 'media' in target:
                        mapping[rel_id] = os.path.basename(target)
    except Exception as e:
        print(f"Error reading relations: {e}")
    return mapping

def docx_to_markdown_with_position(docx_path: str, image_mapping: dict) -> str:
    """将 docx 转换为 Markdown，图片保持在原位置"""
    doc = Document(docx_path)
    markdown_lines = []
    image_counter = 0
    uploaded_images = {}  # 缓存已上传的图片 URL
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # 先检查段落中是否有图片
        paragraph_images = get_paragraph_images(paragraph)
        
        # 如果有图片，先处理图片
        for img_rel_id in paragraph_images:
            if img_rel_id in image_mapping:
                img_filename = image_mapping[img_rel_id]
                image_counter += 1
                # 检查是否已上传
                if img_filename not in uploaded_images:
                    # 需要从原始文件中提取图片数据
                    # 这里简化处理，在实际使用时需要从全局 images dict 获取
                    uploaded_images[img_filename] = f"[图片{image_counter}: {img_filename}]"
                markdown_lines.append(f"![图片{image_counter}]({uploaded_images[img_filename]})")
        
        # 处理段落文本
        if text:
            # 检查是否是标题
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name:
                    level = 1
                    if '1' in style_name:
                        level = 1
                    elif '2' in style_name:
                        level = 2
                    elif '3' in style_name:
                        level = 3
                    elif '4' in style_name:
                        level = 4
                    elif '5' in style_name:
                        level = 5
                    markdown_lines.append(f"{'#' * level} {text}")
                    continue
            
            # 处理粗体和斜体
            formatted_text = text
            for run in paragraph.runs:
                if run.bold and run.italic:
                    formatted_text = formatted_text.replace(run.text, f"***{run.text}***")
                elif run.bold:
                    formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                elif run.italic:
                    formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
            
            markdown_lines.append(formatted_text)
    
    # 处理表格
    for table in doc.tables:
        if len(table.rows) > 0:
            markdown_lines.append("")
            # 表头
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("|" + "|".join([" --- " for _ in header_cells]) + "|")
            # 数据行
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |")
            markdown_lines.append("")
    
    return "\n\n".join(markdown_lines) if markdown_lines else "# 转换结果\n\n文档内容为空"

def docx_to_markdown_enhanced(docx_path: str, images_dict: dict, image_mapping: dict) -> str:
    """增强版转换，图片在原位置显示"""
    doc = Document(docx_path)
    markdown_lines = []
    image_counter = 0
    uploaded_urls = {}
    
    for paragraph in doc.paragraphs:
        line_parts = []
        
        # 处理段落中的每个 run
        for run in paragraph.runs:
            # 检查 run 中是否有图片
            has_image = False
            try:
                blips = run._element.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed_id = blip.get(qn('r:embed'))
                    if embed_id and embed_id in image_mapping:
                        has_image = True
                        img_filename = image_mapping[embed_id]
                        image_counter += 1
                        
                        # 获取图片 URL
                        if img_filename in images_dict:
                            if img_filename not in uploaded_urls:
                                try:
                                    cdn_url = upload_image_to_github(images_dict[img_filename], img_filename)
                                    uploaded_urls[img_filename] = cdn_url
                                    print(f"图片上传成功: {img_filename}")
                                except Exception as e:
                                    uploaded_urls[img_filename] = f"[图片上传失败: {img_filename}]"
                                    print(f"图片上传失败: {e}")
                            
                            line_parts.append(f"![图片{image_counter}]({uploaded_urls[img_filename]})")
            except Exception as e:
                print(f"Error processing run: {e}")
            
            # 处理文本
            if not has_image and run.text:
                text = run.text
                # 处理粗体和斜体
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"
                line_parts.append(text)
        
        # 组合段落内容
        if line_parts:
            combined = ''.join(line_parts)
            if combined.strip():
                # 检查是否是标题
                if paragraph.style and paragraph.style.name:
                    style_name = paragraph.style.name.lower()
                    if 'heading' in style_name:
                        level = 1
                        if '1' in style_name:
                            level = 1
                        elif '2' in style_name:
                            level = 2
                        elif '3' in style_name:
                            level = 3
                        elif '4' in style_name:
                            level = 4
                        elif '5' in style_name:
                            level = 5
                        markdown_lines.append(f"{'#' * level} {combined}")
                        continue
                
                markdown_lines.append(combined)
    
    # 处理表格（表格保持原位置）
    for table in doc.tables:
        if len(table.rows) > 0:
            markdown_lines.append("")
            # 表头
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("|" + "|".join([" --- " for _ in header_cells]) + "|")
            # 数据行
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |")
            markdown_lines.append("")
    
    result = "\n\n".join(markdown_lines) if markdown_lines else "# 转换结果\n\n文档内容为空"
    
    # 如果还有未处理的图片（作为后备），添加到末尾
    if image_counter == 0 and images_dict:
        result += "\n\n## 📷 图片附件\n\n"
        for img_name in images_dict.keys():
            try:
                cdn_url = upload_image_to_github(images_dict[img_name], img_name)
                result += f"![{img_name}]({cdn_url})\n\n"
            except Exception as e:
                result += f"\n*图片 {img_name} 上传失败: {str(e)}*\n\n"
    
    return result

# ========== API 端点 ==========
@app.get("/")
async def root():
    """返回前端页面"""
    possible_paths = [
        Path(__file__).parent / "static" / "index.html",
        Path("/app/static/index.html"),
        Path("static/index.html"),
    ]
    
    for path in possible_paths:
        if path.exists():
            print(f"找到静态文件: {path}")
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                return HTMLResponse(content=content)
            except Exception as e:
                print(f"读取文件失败: {e}")
                continue
    
    return HTMLResponse(content="""
    <html>
        <body style="font-family: sans-serif; text-align: center; padding: 50px;">
            <h1>📄 Word to Markdown Converter</h1>
            <p>服务运行正常，但前端页面未找到。</p>
            <p>请使用 POST /convert 接口上传文件。</p>
            <p><a href="/health">健康检查</a></p>
        </body>
    </html>
    """)

@app.post("/convert")
async def convert_docx(file: UploadFile = File(...)):
    """转换 Word 文档为 Markdown - 图片和表格在原位置"""
    print(f"收到文件: {file.filename}")
    
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="只支持 .docx 格式的文件")
    
    if file.size and file.size > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 10MB")
    
    if not GITHUB_TOKEN or not GITHUB_REPO:
        raise HTTPException(status_code=500, detail="服务器配置错误")
    
    tmpdir = tempfile.mkdtemp()
    
    try:
        input_path = os.path.join(tmpdir, file.filename)
        
        # 保存上传的文件
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)
        print(f"文件保存成功: {input_path}, 大小: {len(content)} 字节")
        
        # 1. 提取所有图片
        images_dict = extract_images_from_docx(input_path)
        print(f"找到 {len(images_dict)} 张图片")
        
        # 2. 获取图片关系映射
        image_mapping = get_image_relation_mapping(input_path)
        print(f"图片关系映射: {image_mapping}")
        
        # 3. 转换为 Markdown（图片在原位置）
        markdown_content = docx_to_markdown_enhanced(input_path, images_dict, image_mapping)
        print(f"转换成功，内容长度: {len(markdown_content)}")
        
        # 4. 确保有内容
        if not markdown_content.strip():
            markdown_content = "# 转换结果\n\n文档内容为空或无法解析"
        
        # 5. 保存结果文件
        output_path = os.path.join(tmpdir, "converted.md")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        print(f"Markdown 文件保存成功: {output_path}")
        print(f"文件大小: {os.path.getsize(output_path)} 字节")
        
        filename = Path(file.filename).stem + ".md"
        
        def cleanup():
            try:
                shutil.rmtree(tmpdir, ignore_errors=True)
                print(f"已清理临时目录: {tmpdir}")
            except Exception as e:
                print(f"清理临时目录失败: {e}")
        
        return FileResponse(
            output_path,
            media_type="text/markdown",
            filename=filename,
            background=BackgroundTask(cleanup)
        )
        
    except Exception as e:
        if tmpdir and os.path.exists(tmpdir):
            shutil.rmtree(tmpdir, ignore_errors=True)
        error_detail = traceback.format_exc()
        print(f"转换失败:\n{error_detail}")
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")

@app.get("/health")
async def health():
    """健康检查端点"""
    return {
        "status": "healthy",
        "github_repo": GITHUB_REPO,
        "github_token_configured": bool(GITHUB_TOKEN),
        "version": "2.1"
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)